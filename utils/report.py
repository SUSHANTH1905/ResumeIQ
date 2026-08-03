"""
report.py
----------
Generates downloadable PDF and DOCX reports summarizing the resume
analysis: candidate details, ATS score breakdown, matched/missing
skills, quality notes, and improvement suggestions.
"""

import io

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
)

from utils.exceptions import ReportGenerationError
from utils.logging_config import get_logger

logger = get_logger("report")


def _safe_text(value, fallback="Not found"):
    return value if value else fallback


def build_pdf_report(parsed_resume: dict, ats_result: dict,
                      skills_found: list, missing_skills: list,
                      suggestions: list, match_percentage: float = None) -> bytes:
    """
    Builds a PDF report in-memory and returns the raw bytes.

    Raises:
        ReportGenerationError: if the PDF cannot be built.
    """
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            topMargin=2 * cm, bottomMargin=2 * cm,
            leftMargin=2 * cm, rightMargin=2 * cm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "TitleCustom", parent=styles["Title"], textColor=colors.HexColor("#2E4053")
        )
        heading_style = ParagraphStyle(
            "HeadingCustom", parent=styles["Heading2"], textColor=colors.HexColor("#2874A6"),
            spaceBefore=14, spaceAfter=6,
        )
        normal = styles["Normal"]

        story = [Paragraph("ResumeIQ - AI Resume Analysis Report", title_style), Spacer(1, 12)]

        story.append(Paragraph("Candidate Information", heading_style))
        info_table_data = [
            ["Name", _safe_text(parsed_resume.get("name"))],
            ["Email", _safe_text(parsed_resume.get("email"))],
            ["Phone", _safe_text(parsed_resume.get("phone"))],
            ["LinkedIn", _safe_text(parsed_resume.get("linkedin"))],
            ["GitHub", _safe_text(parsed_resume.get("github"))],
            ["Projects Detected", str(parsed_resume.get("project_count", 0))],
        ]
        info_table = Table(info_table_data, colWidths=[4 * cm, 11 * cm])
        info_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#D6EAF8")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(info_table)

        story.append(Paragraph("ATS Score Breakdown", heading_style))
        breakdown = ats_result.get("breakdown", {})
        max_scores = ats_result.get("max_scores", {})
        score_data = [["Category", "Score", "Max"]]
        for category, score in breakdown.items():
            score_data.append([category, str(score), str(max_scores.get(category, ""))])
        score_data.append(["TOTAL", str(ats_result.get("total", 0)), "100"])

        score_table = Table(score_data, colWidths=[7 * cm, 4 * cm, 4 * cm])
        score_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2874A6")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#D6EAF8")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ]))
        story.append(score_table)

        if match_percentage is not None:
            story.append(Spacer(1, 8))
            story.append(Paragraph(f"<b>Job Description Match:</b> {match_percentage}%", normal))

        story.append(Paragraph("Skills Found", heading_style))
        story.append(Paragraph(", ".join(skills_found) if skills_found else "No known skills detected.", normal))

        story.append(Paragraph("Missing Skills (vs Job Description)", heading_style))
        if missing_skills:
            story.append(Paragraph(", ".join(missing_skills), normal))
        else:
            story.append(Paragraph("None — great match, or no job description was provided.", normal))

        story.append(Paragraph("Suggestions for Improvement", heading_style))
        bullet_items = [ListItem(Paragraph(s, normal)) for s in (suggestions or [])]
        if bullet_items:
            story.append(ListFlowable(bullet_items, bulletType="bullet"))
        else:
            story.append(Paragraph("No suggestions — resume looks strong!", normal))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as exc:
        logger.exception("Failed to build PDF report")
        raise ReportGenerationError(f"Could not generate PDF report: {exc}") from exc


def build_docx_report(parsed_resume: dict, ats_result: dict,
                       skills_found: list, missing_skills: list,
                       suggestions: list, match_percentage: float = None) -> bytes:
    """
    Builds a DOCX report in-memory and returns the raw bytes.

    Raises:
        ReportGenerationError: if the DOCX cannot be built.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        document = Document()

        title = document.add_heading("ResumeIQ - AI Resume Analysis Report", level=0)
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x28, 0x74, 0xA6)

        document.add_heading("Candidate Information", level=1)
        table = document.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        fields = [
            ("Name", _safe_text(parsed_resume.get("name"))),
            ("Email", _safe_text(parsed_resume.get("email"))),
            ("Phone", _safe_text(parsed_resume.get("phone"))),
            ("LinkedIn", _safe_text(parsed_resume.get("linkedin"))),
            ("GitHub", _safe_text(parsed_resume.get("github"))),
            ("Projects Detected", str(parsed_resume.get("project_count", 0))),
        ]
        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value

        document.add_heading("ATS Score Breakdown", level=1)
        breakdown = ats_result.get("breakdown", {})
        max_scores = ats_result.get("max_scores", {})
        score_table = document.add_table(rows=1, cols=3)
        score_table.style = "Light Grid Accent 1"
        hdr = score_table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Category", "Score", "Max"
        for category, score in breakdown.items():
            row = score_table.add_row().cells
            row[0].text = category
            row[1].text = str(score)
            row[2].text = str(max_scores.get(category, ""))
        total_row = score_table.add_row().cells
        total_row[0].text, total_row[1].text, total_row[2].text = (
            "TOTAL", str(ats_result.get("total", 0)), "100"
        )

        if match_percentage is not None:
            p = document.add_paragraph()
            run = p.add_run(f"Job Description Match: {match_percentage}%")
            run.bold = True
            run.font.size = Pt(11)

        document.add_heading("Skills Found", level=1)
        document.add_paragraph(", ".join(skills_found) if skills_found else "No known skills detected.")

        document.add_heading("Missing Skills (vs Job Description)", level=1)
        document.add_paragraph(
            ", ".join(missing_skills) if missing_skills
            else "None — great match, or no job description was provided."
        )

        document.add_heading("Suggestions for Improvement", level=1)
        if suggestions:
            for s in suggestions:
                document.add_paragraph(s, style="List Bullet")
        else:
            document.add_paragraph("No suggestions — resume looks strong!")

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as exc:
        logger.exception("Failed to build DOCX report")
        raise ReportGenerationError(f"Could not generate DOCX report: {exc}") from exc
